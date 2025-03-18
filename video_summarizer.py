import os
from moviepy import VideoFileClip
import speech_recognition as sr
from transformers import pipeline
from docx import Document
from docx.shared import Inches
from PIL import Image

# Function to handle video file processing
def handle_video_file(video_path):
    audio_file = "extracted_audio.wav"

    # Get base filename without extension
    video_base_name = os.path.splitext(os.path.basename(video_path))[0]

    # Step 1: Extract Audio from the Video
    video = VideoFileClip(video_path)
    video.audio.write_audiofile(audio_file, codec="pcm_s16le")  # Save audio as WAV

    # Step 2: Convert Speech to Text with Timestamps
    speech_recognizer = sr.Recognizer()
    audio_segment_duration = 60  # Process audio in 60-second segments

    with sr.AudioFile(audio_file) as audio_source:
        total_duration = int(video.duration)  # Get the video's total duration
        transcription = []
        annotated_transcriptions = []  # Store speech-marked transcriptions
        time_stamps = []  # Store timestamps for segments

        for start in range(0, total_duration, audio_segment_duration):
            audio_segment = speech_recognizer.record(audio_source, duration=audio_segment_duration)
            timestamp = f"{start // 3600:02}:{(start % 3600) // 60:02}:{start % 60:02}"
            time_stamps.append(timestamp)  # Record timestamp for segment

            try:
                transcribed_text = speech_recognizer.recognize_google(audio_segment, language="en-US")
                transcription.append(f"[{timestamp}] {transcribed_text}")
                annotated_transcriptions.append(f"Speech segment for [{timestamp}]:\n- {transcribed_text}")
            except sr.UnknownValueError:
                transcription.append(f"[{timestamp}] [Speech Unclear]")
                annotated_transcriptions.append(f"Speech segment for [{timestamp}]:\n- [Speech Unclear]")
            except sr.RequestError:
                transcription.append(f"[{timestamp}] [API Error]")
                annotated_transcriptions.append(f"Speech segment for [{timestamp}]:\n- [API Error]")

    # Step 3: Summarize Speech Transcriptions Using BART Model
    summarizer = pipeline("summarization", model="facebook/bart-large-xsum")

    # Join the transcriptions into a string
    concatenated_transcriptions = "\n\n".join(annotated_transcriptions)

    # Split into paragraphs
    transcribed_paragraphs = concatenated_transcriptions.split("\n\n")

    # Initialize list for summaries
    summary_texts = []

    # Generate summaries for each paragraph with corresponding timestamp
    for index, (paragraph, timestamp) in enumerate(zip(transcribed_paragraphs, time_stamps)):
        if paragraph.strip():  # Skip empty paragraphs
            summary = summarizer(paragraph, max_length=20, min_length=10, do_sample=False)
            summary_texts.append(f"Summary of Paragraph {index+1} [{timestamp}]:\n{summary[0]['summary_text']}")
        else:
            summary_texts.append(f"Summary of [{timestamp}]:\n[Empty Paragraph]")

    # Step 4: Save Summaries and Frames to a Word Document
    document = Document()
    document.add_heading(f"Summary of video: {video_base_name}", 0)

    for index, summary in enumerate(summary_texts):
        document.add_paragraph(summary)

        # Step 5: Extract and Add Frame Below Summary
        segment_start_time = index * audio_segment_duration  # Calculate start time for current segment
        image_filename = f"frame_{index}.png"  # Temporary image file name

        # Extract frame from video
        frame_image = video.get_frame(segment_start_time)
        frame = Image.fromarray(frame_image)
        frame.save(image_filename)

        # Add image to Word document
        document.add_picture(image_filename, width=Inches(3))

        # Delete temporary image file after adding it to document
        os.remove(image_filename)

    # Save the Word document with summaries and frames
    doc_output_filename = f"{video_base_name}_summarized_with_images.docx"
    document.save(doc_output_filename)
    print(f"Summaries and frames saved in {doc_output_filename}")

    # Delete temporary audio file
    if os.path.exists(audio_file):
        os.remove(audio_file)
        print(f"Deleted temporary audio file: {audio_file}")
    else:
        print(f"{audio_file} not found.")

# Function to process all videos in a folder
def process_video_folder(folder_path):
    video_files = [file for file in os.listdir(folder_path) if file.lower().endswith(('.mp4', '.avi', '.mov', '.mkv'))]

    if not video_files:
        print("No video files found in the specified folder.")
        return

    for video_filename in video_files:
        video_filepath = os.path.join(folder_path, video_filename)
        print(f"Processing video: {video_filename}")
        handle_video_file(video_filepath)

# Main function for user interaction
def main():
    while True:
        user_choice = input(
            "Enter 'file' to process a single video file, 'folder' for multiple videos, or 'quit' to exit: "
        ).strip().lower()

        if user_choice in ['file', 'folder', 'quit']:
            break
        else:
            print("Invalid choice! Please enter 'file', 'folder', or 'quit'.")

        if user_choice == 'file':
            while True:  # Handle incorrect file path inputs
                video_file_path = input("Enter the video file path: ").strip()
                if os.path.exists(video_file_path):
                    handle_video_file(video_file_path)
                    break  # Exit loop after successful processing
                else:
                    print(f"Invalid file path: {video_file_path}. Please provide a valid file.")

        elif user_choice == 'folder':
            while True:  # Handle incorrect folder path inputs
                folder_path = input("Enter the folder path with video files: ").strip()
                if os.path.exists(folder_path) and os.path.isdir(folder_path):
                    process_video_folder(folder_path)
                    break  # Exit loop after successful processing
                else:
                    print(f"Invalid folder path: {folder_path}. Please provide a valid folder.")

        elif user_choice == 'quit':
            print("Exiting program.")
            break  # Exit the main loop

# Run the script if executed directly
if __name__ == "__main__":
    main()
